# Nombre del reporte
# Findins con titulo y recursos afectados
# En base al titulo se obtienen de una DB  (so far json file)
# Parametros a recibir:
#   Idioma: eng o spa
#   Tipo Reporte: S o F (indica si es Estado o Final)
#   json con los titulos de los findings a buscar
import datetime,docx,json,os,ast,base64,time

import gc

from copy import deepcopy
from docx.shared import Inches
from .. import constants
from ..mongo import mongo
from io import BytesIO
from PIL import Image, ImageOps

# Variables Generales
ROOT_DIR = os.path.dirname(os.path.abspath(__file__))
# Determina el idioma del reporte
eng = False
# Determina el tipo de reporte
estado = True
# Estos indices son para copiar el finding de inicio a fin varian segundo el tipo de reporte
indexNros = (0, 0, 0)
missing_findings=""
# Documento
doc = docx.Document()


############################################################
# PARA REPORTES ESTADO DE AVANCE EN ESPAÑOL O FINAL SLATAM #
############################################################

# Buscamos la nota en el listado de los paragraph
def agregarNota(p, nota):
    if "Nota" in p.text:
        p.text = "Nota: " + nota
    elif "Note:" in p.text:
        p.text = "Note: " + nota

def add_screenshot(jsonFinding):
    path=""
    if jsonFinding['image_string'] != None:
        buf = BytesIO(base64.b64decode(jsonFinding['image_string']))
        img = Image.open(buf)
        ImageOps.expand(img,border=10,fill='black').save(ROOT_DIR+"/image.png")
        img.save(ROOT_DIR+"/image.png")
        path = ROOT_DIR+"/image.png"
    return path

def delete_screenshot(image_path):
    os.remove(image_path)

def add_cves(jsonFinding):
    message=""
    if jsonFinding['TITLE'] == constants.OUTDATED_3RD_LIBRARIES['spanish_name'] or jsonFinding['TITLE'] == constants.OUTDATED_3RD_LIBRARIES['english_name']:
        if jsonFinding['extra_info'] != None :
            message = jsonFinding['extra_info']
    return message

# Agregar info en celda sin romper formato
def addFindingInfo(table, language, urls):
    # ---------------------Recursos afectados--------------------------------
    paragNumber = 0
    resourcesTitle = table.cell(0, 0).paragraphs[paragNumber]
    # Cargamos los recursos afectados y borramos los valores que trae el documento por defecto
    tableResources = table.cell(1, 0)
    paragraph = tableResources.paragraphs[paragNumber]
    urlsSize = len(urls)
    if urlsSize > 1:
        resourcesTitle.text = constants.recursoAfectadoPlu_EN if language else constants.recursoAfectadoPlu_ES
        if urlsSize < 3:
            for rec in urls:
                p = deepcopy(paragraph)
                p.text = rec
                paragraph._p.addnext(p._p)
                paragNumber += 1
        else:
            p = deepcopy(paragraph)
            p.text = "CREAR APENDICE"
            paragraph._p.addnext(p._p)
            paragNumber += 1
        # Borro primer recurso ya que no se necesita
        p1 = tableResources.paragraphs[0]._element
        p1.getparent().remove(p1)
        p1._p = p1._element = None
    else:
        # Al ser un solo recurso va en singular
        resourcesTitle.text = constants.recursoAfectadoSin_EN if language else constants.recursoAfectadoSin_ES
        paragraph.text = urls[0]


def clonarTemplateYAgregarFinding(doc, indexNros, language, jsonFinding):
    nroParagraph = indexNros[0]
    nroTable = indexNros[1]
    nroPageBreak = indexNros[2]
    urls = jsonFinding['resourceAf']
    refUrls = jsonFinding['RECOMMENDATION']['URLS']
    page_break = doc.paragraphs[nroPageBreak]
    templateParag = deepcopy(doc.paragraphs[nroParagraph])
    # Titulo finding
    runNro = 9
    if not indexNros[3]: runNro = 7
    templateParag.runs[runNro].text = jsonFinding['TITLE']
    # Agarra pagebreak anterior y pone el nuevo titulo abajo
    page_break._p.addnext(templateParag._p)
    # Agregar todo el contenido
    new_tbl = deepcopy(doc.tables[nroTable])
    p, pb = templateParag._p, page_break._p
    new_pb = deepcopy(pb)
    p.addnext(new_pb)
    # Clonamos toda la informacion que es parte del finding (Observacion, tabla, impacto, etc.)
    #Excepto el parrafo de la imagen ya que si lo clonamos se puede romper el formato
    obsTitle = deepcopy(doc.paragraphs[nroParagraph + 2])
    obsText = deepcopy(doc.paragraphs[nroParagraph + 3])
    urlTitle = deepcopy(doc.paragraphs[nroParagraph + 4])
    obsNote = deepcopy(doc.paragraphs[nroParagraph + 7])
    legendTable = deepcopy(doc.paragraphs[nroParagraph + 8])
    legendFigure = deepcopy(doc.paragraphs[nroParagraph + 10])
    impactTitle = deepcopy(doc.paragraphs[nroParagraph + 11])
    impactDesc = deepcopy(doc.paragraphs[nroParagraph + 12])
    likelihoodTitle = deepcopy(doc.paragraphs[nroParagraph + 13])
    likelihoodDesc = deepcopy(doc.paragraphs[nroParagraph + 14])
    recomendTitle = deepcopy(doc.paragraphs[nroParagraph + 15])
    recomendNote = deepcopy(doc.paragraphs[nroParagraph + 16])
    recomendText = deepcopy(doc.paragraphs[nroParagraph + 17])
    referenceTitle = deepcopy(doc.paragraphs[nroParagraph + 18])
    referenceLinkLegend = deepcopy(doc.paragraphs[nroParagraph + 19])
    # Cargamos la informacion en base a lo que se tenga
    # URLS RECOMENDACIONES
    if len(refUrls) >= 1:
        for url in refUrls:
            urlRef = deepcopy(doc.paragraphs[nroParagraph + 20])
            urlRef.text = url
            p.addnext(urlRef._p)
        if len(refUrls) > 1:
            referenceLinkLegend.text = constants.enlacesRecomendacion_EN if language else constants.enlacesRecomendacion_ES
            p.addnext(referenceLinkLegend._p)
        else:
            p.addnext(referenceLinkLegend._p)
        p.addnext(referenceTitle._p)
    recomendText.text = jsonFinding['RECOMMENDATION']['TITLE']
    p.addnext(recomendText._p)
    p.addnext(recomendNote._p)
    p.addnext(recomendTitle._p)
    p.addnext(likelihoodDesc._p)
    p.addnext(likelihoodTitle._p)
    impactDesc.text = jsonFinding['IMPLICATION']
    p.addnext(impactDesc._p)
    p.addnext(impactTitle._p)
    p.addnext(legendFigure._p)
    #Agregamos imagen si es que hay
    figureText = doc.paragraphs[nroParagraph + 9]
    figureText.clear()
    figureText.add_run()
    image_path = add_screenshot(jsonFinding)
    if image_path:
        figureText.runs[0].add_picture(image_path,width=Inches(5.33), height=Inches(4.0))
        delete_screenshot(image_path)
    para = deepcopy(figureText)
    p.addnext(para._p)
    p.addnext(legendTable._p)
    mid_tbl = deepcopy(doc.tables[nroTable + 1])
    p.addnext(mid_tbl._tbl)
    if jsonFinding['OBSERVATION']['NOTE']:
        agregarNota(obsNote, jsonFinding['OBSERVATION']['NOTE'])
    else:
        obsNote.text = ""
    p.addnext(obsNote._p)
    for url in urls:
        urlExample = deepcopy(doc.paragraphs[nroParagraph + 5])
        urlExample.text = url
        p.addnext(urlExample._p)
    if len(urls) == 1:
        urlTitle.text = constants.urlAfectada_EN if language else constants.urlAfectada_ES
        p.addnext(urlTitle._p)
    else:
        p.addnext(urlTitle._p)
    obsText.text = jsonFinding['OBSERVATION']['TITLE']
    obsText.text+= add_cves(jsonFinding)
    p.addnext(obsText._p)
    p.addnext(obsTitle._p)
    # Espacio vacio para respetar el formato
    p.addnext(deepcopy(doc.paragraphs[nroParagraph + 1])._p)
    p.addnext(new_tbl._tbl)
    addFindingInfo(new_tbl, language, urls)


def crearReporte(language, reportType, findings):
    eng = True if language == 'eng' else False
    estado = True if reportType == 'S' else False
    global doc,missing_findings
    # Template a utilizar acorde al tipo de reporte que se necesite generar, va a variar dependiendo del idioma, avance o final
    if not eng:
        if estado:
            doc = docx.Document(ROOT_DIR + '/out/TEMPLATE - REPORTE DE ESTADO - DELOITTE S-LATAM - v1.1.docm')
            indexNros = (13, 3, 34, estado)
        else:
            doc = docx.Document(ROOT_DIR + '/out/TEMPLATE - REPORTE FINAL - DELOITTE S-LATAM - v1.docm')
            indexNros = (76, 5, 97, estado)
    else:
        if estado:
            # Voy a necesitar los reportes para ingles POR AHORA USA REPORTE DE ESTADO
            doc = docx.Document(ROOT_DIR + '/out/TEMPLATE - REPORTE DE ESTADO - DELOITTE S-LATAM - v1.1.docm')
            indexNros = (13, 3, 34, estado)
        else:
            doc = docx.Document(ROOT_DIR + '/out/TEMPLATE - REPORTE FINAL - DELOITTE S-LATAM - v1.docm')
            indexNros = (76, 5, 97, estado)

    # Itero sobre la lista recibida si el finding existe lo agrego sino pongo un mensaje de alerta
    print("Generando Reporte espere")
    for finding in findings:
        # Obtenemos finding de la KB
        json_value = mongo.get_specific_finding_info(finding, language)
        if json_value:
            clonarTemplateYAgregarFinding(doc, indexNros, eng, json_value)
        else:
            #TODO mensaje a slack
            print("The following finding was not found: "+finding['title'])
            missing_findings += finding['title']+'\n'
        gc.collect()
    print("Se genero el reporte con los findings que fueron encontrados")
    d1 = datetime.datetime.now().strftime('%Y%m%d')
    name = ""
    # Muy probable que esto cambie
    if not eng:
        if estado:
            name += "REPORTE_DE_ESTADO_" + d1 + ".docm"
            doc.save(ROOT_DIR + '/out/' + name)
        else:
            name += "REPORTE_FINAL_" + d1 + ".docm"
            doc.save(ROOT_DIR + '/out/' + name)
    else:
        name += "REPORTE_CON_FINDINGS_INGLES-" + d1 + ".docm"
        doc.save(ROOT_DIR + '/out/' + name)
    gc.collect()
    path = ROOT_DIR + '/out/' + name
    return path,missing_findings